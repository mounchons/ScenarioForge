"""
Turns the real client template docs/template/F-DA-003_FunctionalSpec.docx into a
docxtpl-ready template (fda003_template.docx) for render_fda003_docx.py.

Run once (or whenever the source .docx or the mapping below changes):
    python patch_fda003_template.py

What it does, section by section (see references/word-export.md for the full rationale):
  - Doc control (Table 0): company/project code/name -> real tags, fed from doc-meta.json.
  - Revision History (Table 2): sample row -> {%tr for/endfor%} loop, fed from doc-meta.json.
  - System Diagram / Sequence Diagram / ER Diagram headings: blank paragraph after each ->
    an image placeholder, fed from the assembled system-design-document.md's mermaid blocks.
  - Database Specification (Heading3 + 4 description lines + its field table): wrapped in a
    whole-group {% for entity in entities %}...{% endfor %} block, one heading+table per
    Data Dictionary entity. DB Name defaults to the entity name (EF-Core-style convention);
    Description/Server UAT/Server Production are left blank -- not derivable from the DD.
  - Database Normalization (Table 9): {%tr %} row-loop, one row per entity, defaulted to
    "New" + 1NF/2NF/3NF checked (methodological inference from domain-design's own modeling
    rules, NOT invented business fact) -- BCNF left unchecked, Remark left blank.
  - Everything else (Screen/Process/Document/Job function-spec tables, all NON-FUNCTIONAL
    infra tables: Server/DB-env/API-env/Security/Performance/Reliability): domain-design does
    not produce this data (infra detail, or -- for Screen -- ui-mockup's own territory), so
    these are left as static tables with the stale literal {{xxx}}/{{Low, Medium, High}}/
    {{number}} placeholders stripped out (so they don't break Jinja parsing) but the row/column
    structure intact, ready for a human to fill in by hand.
"""
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parents[7]
SRC = REPO_ROOT / "docs" / "template" / "F-DA-003_FunctionalSpec.docx"
OUT = Path(__file__).parent / "fda003_template.docx"

PLACEHOLDER_RE = re.compile(r"\{\{[^}]*\}\}")


def strip_placeholders(text):
    """Remove stale literal {{...}} markers from non-derivable fields, keeping any literal
    prefix (e.g. 'F_SCR_{{number}}' -> 'F_SCR_')."""
    return PLACEHOLDER_RE.sub("", text)


def clean_table_cells(table, skip_rows=0):
    for row in table.rows[skip_rows:]:
        for cell in row.cells:
            if "{{" in cell.text:
                cell.text = strip_placeholders(cell.text)


def find_paragraph(doc, text, style=None, start_after=None):
    """Return the first paragraph matching text (and style if given), searching in document
    order. If start_after is a paragraph, only consider paragraphs after it."""
    found_start = start_after is None
    for p in doc.paragraphs:
        if not found_start:
            if p._p is start_after._p:
                found_start = True
            continue
        if p.text.strip() == text and (style is None or (p.style and p.style.name == style)):
            return p
    raise ValueError(f"paragraph not found: {text!r} (style={style})")


def find_table_after(doc, paragraph):
    """Return the table immediately following `paragraph` in document order."""
    body = doc.element.body
    started = False
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    for child in body.iterchildren():
        if child is paragraph._p:
            started = True
            continue
        if started:
            if isinstance(child, CT_Tbl):
                return Table(child, doc)
            # keep scanning past blank paragraphs until we hit the table
    raise ValueError("no table found after paragraph")


def add_row_before(table, ref_row):
    new_row = table.add_row()
    ref_row._tr.addprevious(new_row._tr)
    return new_row


def delete_row(row):
    row._tr.getparent().remove(row._tr)


def wrap_row_as_tr_loop(table, data_row_index, for_tag, endfor_tag="{%tr endfor %}"):
    """Turn table.rows[data_row_index] into the repeating row of a {%tr %} loop: insert a
    FOR-marker row immediately before it and an ENDFOR-marker row immediately after it, then
    delete every other row that follows (the template's blank filler rows). Must capture the
    "row after data" reference BEFORE inserting the FOR-row, since inserting shifts indices --
    grabbing it late (e.g. via a freshly re-read table.rows[data_row_index + 1]) silently lands
    on the data row itself instead (bit us once already; see git history of this file)."""
    data_row = table.rows[data_row_index]
    row_after_data = table.rows[data_row_index + 1] if len(table.rows) > data_row_index + 1 else None

    for_row = add_row_before(table, data_row)
    for_row.cells[0].text = for_tag

    if row_after_data is not None:
        endfor_row = add_row_before(table, row_after_data)
    else:
        endfor_row = table.add_row()
    endfor_row.cells[0].text = endfor_tag

    # rows after the endfor marker are the template's original blank filler rows -- drop them
    keep_through = data_row_index + 3  # header(s)... for-row, data-row, endfor-row
    for row in list(table.rows[keep_through:]):
        delete_row(row)

    return data_row


def main():
    doc = Document(str(SRC))

    # ---------- Table 0: doc control ----------
    t0 = doc.tables[0]
    t0.rows[0].cells[1].text = "{{ client_company_name }}"
    t0.rows[1].cells[1].text = "{{ project_code }}"
    t0.rows[2].cells[1].text = "{{ project_name }}"

    # ---------- Table 2: revision history ----------
    t2 = doc.tables[2]
    t2.rows[1].cells[0].text = "{{ rev.version }}"
    t2.rows[1].cells[1].text = "{{ rev.date }}"
    t2.rows[1].cells[2].text = "{{ rev.author }}"
    t2.rows[1].cells[3].text = "{{ rev.description }}"
    wrap_row_as_tr_loop(t2, 1, "{%tr for rev in revisions %}")

    # ---------- System / Sequence / ER diagram images ----------
    sys_heading = find_paragraph(doc, "System Diagram", style="Heading 2")
    seq_heading = find_paragraph(doc, "Sequence Diagram", style="Heading 2")
    er_heading = find_paragraph(doc, "E-R DIAGRAM", style="Heading 2")

    def next_blank_paragraph(after):
        paragraphs = doc.paragraphs
        idx = next(i for i, p in enumerate(paragraphs) if p._p is after._p)
        for cand in paragraphs[idx + 1:]:
            if cand.style and cand.style.name.startswith("Heading"):
                raise ValueError(f"no blank paragraph found after {after.text!r} before next heading")
            if not cand.text.strip():
                return cand
        raise ValueError("ran out of paragraphs")

    next_blank_paragraph(sys_heading).text = "{{ s_system_diagram_img }}"
    next_blank_paragraph(seq_heading).text = "{{ s_sequence_diagram_img }}"
    next_blank_paragraph(er_heading).text = "{{ s_er_diagram_img }}"

    # ---------- Database Specification: whole-group loop per entity ----------
    dbspec_heading3 = find_paragraph(doc, "{{Table Name, View, Store Procedure, SQL Function}}", style="Heading 3")
    all_paragraphs = doc.paragraphs
    idx = next(i for i, p in enumerate(all_paragraphs) if p._p is dbspec_heading3._p)
    desc_p, dbname_p, uat_p, prod_p = all_paragraphs[idx + 1: idx + 5]
    assert desc_p.text.startswith("Description:"), desc_p.text
    assert dbname_p.text.startswith("DB Name:"), dbname_p.text
    assert uat_p.text.startswith("Server Name UAT:"), uat_p.text
    assert prod_p.text.startswith("Server Name Production:"), prod_p.text
    t8 = find_table_after(doc, prod_p)
    assert t8.rows[0].cells[0].text.strip() == "Field Name", t8.rows[0].cells[0].text

    dbspec_heading3.insert_paragraph_before("{% for entity in entities %}")
    dbspec_heading3.text = "{{ entity.name }}"
    desc_p.text = "Description: "
    dbname_p.text = "DB Name: {{ entity.name }}"
    uat_p.text = "Server Name UAT: "
    prod_p.text = "Server Name Production: "

    t8.rows[1].cells[0].text = "{{ field.field }}"
    t8.rows[1].cells[1].text = "{{ field.type }}"
    t8.rows[1].cells[2].text = "{{ field.description }}"
    t8.rows[1].cells[3].text = ""
    wrap_row_as_tr_loop(t8, 1, "{%tr for field in entity.fields %}")

    # the endfor for the whole-group loop goes right before the next heading
    dbnorm_heading = find_paragraph(doc, "DATABASE NORMALIZATION", style="Heading 2", start_after=prod_p)
    dbnorm_heading.insert_paragraph_before("{% endfor %}")

    # ---------- Database Normalization: row-loop per entity ----------
    t9 = doc.tables[9]
    t9.rows[1].cells[0].text = "{{ loop.index }}"
    t9.rows[1].cells[1].text = "{{ entity.name }}"
    t9.rows[1].cells[2].text = "New"
    t9.rows[1].cells[3].text = "☐"       # Unnormalize - unchecked
    t9.rows[1].cells[4].text = "☑"       # 1NF - checked (domain-design's own modeling convention)
    t9.rows[1].cells[5].text = "☑"       # 2NF - checked
    t9.rows[1].cells[6].text = "☑"       # 3NF - checked
    t9.rows[1].cells[7].text = "☐"       # BCNF - left for manual review
    t9.rows[1].cells[8].text = ""
    wrap_row_as_tr_loop(t9, 1, "{%tr for entity in entities %}")

    # ---------- Clean stale placeholders in non-derivable tables (structure kept, manual fill) ----------
    for ti in (4, 5, 6, 7, 11, 12, 13, 14, 15):
        clean_table_cells(doc.tables[ti], skip_rows=0)
    # cached ToC field result for the (now looped) Database Specification heading -- Word
    # regenerates this text from the real headings once the user updates the field, but the
    # stale cached text has commas/spaces that are invalid Jinja, so neutralize it too
    for p in doc.paragraphs:
        if p.text.strip().startswith("3.5.1") and "{{Table Name" in p.text:
            p.text = "3.5.1\tDatabase Specification\t7"
    # fix the mislabeled NF_SEQ id prefix under the SECURITY heading (cosmetic copy-paste leftover)
    for row in doc.tables[13].rows:
        for cell in row.cells:
            if "NF_SEQ" in cell.text:
                cell.text = cell.text.replace("NF_SEQ", "NF_SEC")

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

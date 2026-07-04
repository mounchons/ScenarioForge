"""
Builds sdd_template.docx — the docxtpl (Jinja2-in-docx) template that
render_sdd_docx.py fills with a project's system-design-document.md content.

Run once (or whenever the template layout needs to change):
    python build_template.py

Regenerating overwrites sdd_template.docx in this folder.
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

THAI_FONT = "TH Sarabun New"
HEADING_COLOR = RGBColor(0x1F, 0x3B, 0x57)

OUT_PATH = os.path.join(os.path.dirname(__file__), "sdd_template.docx")


def set_thai_font(run, size=16, bold=False, color=None):
    run.font.name = THAI_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), THAI_FONT)
    rFonts.set(qn('w:cs'), THAI_FONT)


def add_field(paragraph, field_code):
    """Insert a raw Word field (e.g. TOC) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "Right-click here and choose 'Update Field' to build the table of contents."
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_separate)
    r_element.append(fld_text)
    r_element.append(fld_end)


def heading(doc, text, level=1, size=None, color=HEADING_COLOR):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    default_sizes = {1: 20, 2: 18, 3: 16}
    set_thai_font(run, size=size or default_sizes.get(level, 16), bold=True, color=color)
    return p


def para(doc, text="", size=16, bold=False, align=None, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_thai_font(run, size=size, bold=bold)
    run.font.italic = italic
    return p


def style_table(table):
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_thai_font(r, size=14)
                if not p.runs:
                    set_thai_font(p.add_run(""), size=14)


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = THAI_FONT
    normal.font.size = Pt(16)
    normal_rPr = normal.element.get_or_add_rPr()
    normal_rFonts = normal_rPr.find(qn('w:rFonts'))
    if normal_rFonts is None:
        normal_rFonts = OxmlElement('w:rFonts')
        normal_rPr.append(normal_rFonts)
    normal_rFonts.set(qn('w:eastAsia'), THAI_FONT)
    normal_rFonts.set(qn('w:cs'), THAI_FONT)

    for section in doc.sections:
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)

    # ---------- Cover page ----------
    for _ in range(4):
        doc.add_paragraph()
    para(doc, "{{ project_name }}", size=32, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "System Design Document", size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "จัดทำเพื่อ (Prepared for): {{ client_name }}", size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "เวอร์ชันเอกสาร (Version): {{ doc_version }}", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "วันที่ (Date): {{ doc_date }}", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "จัดทำโดย (Prepared by): {{ author }}", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------- Revision history ----------
    # docxtpl's {%tr %} marker consumes the WHOLE <w:tr> it sits in and collapses it to a
    # plain {% %} tag -- so the for/endfor markers must live in their OWN rows, separate
    # from the data row that actually repeats (see references/word-export.md for detail).
    heading(doc, "ประวัติการแก้ไขเอกสาร (Revision History)", level=1)
    table = doc.add_table(rows=2, cols=4)
    style_table(table)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["เวอร์ชัน", "วันที่", "ผู้จัดทำ", "รายละเอียดการแก้ไข"]):
        cell.text = ""
        set_thai_font(cell.paragraphs[0].add_run(text), size=14, bold=True)
    for_row = table.rows[1].cells
    for_row[0].text = ""
    set_thai_font(for_row[0].paragraphs[0].add_run("{%tr for rev in revisions %}"), size=14)
    data_row = table.add_row().cells
    data_row[0].text = ""
    set_thai_font(data_row[0].paragraphs[0].add_run("{{ rev.version }}"), size=14)
    data_row[1].text = ""
    set_thai_font(data_row[1].paragraphs[0].add_run("{{ rev.date }}"), size=14)
    data_row[2].text = ""
    set_thai_font(data_row[2].paragraphs[0].add_run("{{ rev.author }}"), size=14)
    data_row[3].text = ""
    set_thai_font(data_row[3].paragraphs[0].add_run("{{ rev.description }}"), size=14)
    end_row = table.add_row().cells
    end_row[0].text = ""
    set_thai_font(end_row[0].paragraphs[0].add_run("{%tr endfor %}"), size=14)
    doc.add_page_break()

    # ---------- Table of contents ----------
    heading(doc, "สารบัญ (Table of Contents)", level=1)
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    # ---------- Section 1: Introduction & Overview ----------
    heading(doc, "1. บทนำและภาพรวมระบบ (Introduction & Overview)", level=1)
    para(doc, "{{p s1_intro }}")
    doc.add_page_break()

    # ---------- Section 2: System Requirements ----------
    heading(doc, "2. ความต้องการของระบบ (System Requirements)", level=1)
    para(doc, "{{p s2_requirements }}")
    doc.add_page_break()

    # ---------- Section 3: Module Overview ----------
    heading(doc, "3. ภาพรวมโมดูล (Module Overview)", level=1)
    para(doc, "{{p s3_modules }}")
    doc.add_page_break()

    # ---------- Section 4: Data Model ----------
    heading(doc, "4. แบบจำลองข้อมูล (Data Model)", level=1)
    para(doc, "{{p s4_datamodel }}")
    doc.add_page_break()

    # ---------- Section 5: Data Flow Diagram ----------
    heading(doc, "5. แผนภาพการไหลของข้อมูล (Data Flow Diagram)", level=1)
    heading(doc, "Level 0 — Context Diagram", level=2)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run("{{ s5_dfd_l0_img }}")
    heading(doc, "Level 1", level=2)
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run("{{ s5_dfd_l1_img }}")
    doc.add_page_break()

    # ---------- Section 6: Flow Diagrams (repeating) ----------
    heading(doc, "6. แผนภาพขั้นตอนการทำงาน (Flow Diagrams)", level=1)
    doc.add_paragraph("{% for flow in s6_flows %}")
    fh = doc.add_heading(level=2)
    set_thai_font(fh.add_run("{{ flow.title }}"), size=18, bold=True)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("{{ flow.image }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_page_break()

    # ---------- Section 7: ER Diagram ----------
    heading(doc, "7. แผนภาพความสัมพันธ์ข้อมูล (ER Diagram)", level=1)
    p_er = doc.add_paragraph()
    p_er.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_er.add_run("{{ s7_er_img }}")
    doc.add_page_break()

    # ---------- Section 8: Data Dictionary ----------
    heading(doc, "8. พจนานุกรมข้อมูล (Data Dictionary)", level=1)
    dd_table = doc.add_table(rows=2, cols=7)
    style_table(dd_table)
    dd_hdr = dd_table.rows[0].cells
    for cell, text in zip(dd_hdr, ["Entity", "Field", "Type", "Null", "Key", "Constraint", "Description"]):
        cell.text = ""
        set_thai_font(cell.paragraphs[0].add_run(text), size=13, bold=True)
    dd_for = dd_table.rows[1].cells
    dd_for[0].text = ""
    set_thai_font(dd_for[0].paragraphs[0].add_run("{%tr for row in dd_rows %}"), size=13)
    dd_body = dd_table.add_row().cells
    dd_body[0].text = ""
    set_thai_font(dd_body[0].paragraphs[0].add_run("{{ row.entity }}"), size=13)
    dd_body[1].text = ""
    set_thai_font(dd_body[1].paragraphs[0].add_run("{{ row.field }}"), size=13)
    dd_body[2].text = ""
    set_thai_font(dd_body[2].paragraphs[0].add_run("{{ row.type }}"), size=13)
    dd_body[3].text = ""
    set_thai_font(dd_body[3].paragraphs[0].add_run("{{ row.null }}"), size=13)
    dd_body[4].text = ""
    set_thai_font(dd_body[4].paragraphs[0].add_run("{{ row.key }}"), size=13)
    dd_body[5].text = ""
    set_thai_font(dd_body[5].paragraphs[0].add_run("{{ row.constraint }}"), size=13)
    dd_body[6].text = ""
    set_thai_font(dd_body[6].paragraphs[0].add_run("{{ row.description }}"), size=13)
    dd_end = dd_table.add_row().cells
    dd_end[0].text = ""
    set_thai_font(dd_end[0].paragraphs[0].add_run("{%tr endfor %}"), size=13)
    doc.add_page_break()

    # ---------- Section 9: Sitemap ----------
    heading(doc, "9. ผังเว็บไซต์ (Sitemap)", level=1)
    p_site = doc.add_paragraph()
    p_site.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_site.add_run("{{ s9_sitemap_img }}")
    doc.add_page_break()

    # ---------- Section 10: User Roles & Permissions ----------
    heading(doc, "10. บทบาทผู้ใช้และสิทธิ์การเข้าถึง (User Roles & Permissions)", level=1)
    role_table = doc.add_table(rows=2, cols=3)
    style_table(role_table)
    role_hdr = role_table.rows[0].cells
    for cell, text in zip(role_hdr, ["Role", "Permissions", "Accessible Pages"]):
        cell.text = ""
        set_thai_font(cell.paragraphs[0].add_run(text), size=14, bold=True)
    role_for = role_table.rows[1].cells
    role_for[0].text = ""
    set_thai_font(role_for[0].paragraphs[0].add_run("{%tr for role in role_rows %}"), size=14)
    role_body = role_table.add_row().cells
    role_body[0].text = ""
    set_thai_font(role_body[0].paragraphs[0].add_run("{{ role.name }}"), size=14)
    role_body[1].text = ""
    set_thai_font(role_body[1].paragraphs[0].add_run("{{ role.permissions }}"), size=14)
    role_body[2].text = ""
    set_thai_font(role_body[2].paragraphs[0].add_run("{{ role.pages }}"), size=14)
    role_end = role_table.add_row().cells
    role_end[0].text = ""
    set_thai_font(role_end[0].paragraphs[0].add_run("{%tr endfor %}"), size=14)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()

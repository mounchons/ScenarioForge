"""
Renders a client-ready Word document from a ScenarioForge system-design-document.md
(the file domain-design's /deliver-docs command assembles) using sdd_template.docx.

Usage:
    python render_sdd_docx.py <path-to-system-design-document.md> <output.docx> [--meta doc-meta.json]

What it does:
  1. Parses the 10-section markdown (## headers, ```mermaid blocks, the Data Dictionary table).
  2. Renders every mermaid block to a PNG via `npx @mermaid-js/mermaid-cli` (mmdc). Requires Node.js;
     no local install needed (npx --yes fetches mermaid-cli on first run).
  3. Fills sdd_template.docx (docxtpl) with the parsed text/tables/images and writes the final .docx.

Diagram images and .mmd intermediates are written next to the output file, under <output-stem>_diagrams/.
"""
import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docxtpl import DocxTemplate, InlineImage, RichText
from docx.shared import Mm

TEMPLATE_PATH = Path(__file__).parent / "sdd_template.docx"

SECTION_TITLES = {
    1: "Introduction & Overview",
    2: "System Requirements",
    3: "Module Overview",
    4: "Data Model",
    5: "Data Flow Diagram",
    6: "Flow Diagrams",
    7: "ER Diagram",
    8: "Data Dictionary",
    9: "Sitemap",
    10: "User Roles & Permissions",
}


def split_sections(md_text):
    """Split the SDD markdown into {section_number: raw_body_text}."""
    pattern = re.compile(
        r"^##\s*(\d{1,2})\.\s*(.+?)\s*$",
        re.MULTILINE,
    )
    matches = list(pattern.finditer(md_text))
    sections = {}
    for i, m in enumerate(matches):
        num = int(m.group(1))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md_text)
        sections[num] = md_text[start:end].strip()
    return sections


def extract_mermaid_blocks(body):
    """Return (prose_without_mermaid, [mermaid_code, ...]) for a section body."""
    blocks = re.findall(r"```mermaid\s*\n(.*?)```", body, re.DOTALL)
    prose = re.sub(r"```mermaid\s*\n.*?```", "", body, flags=re.DOTALL).strip()
    return prose, [b.strip() for b in blocks]


def extract_subheadings(body):
    """Split a section body on '### Title' markers -> [(title, text), ...]. Falls back to
    a single ('', body) entry when there are no ### subheadings (used for Flow Diagrams)."""
    pattern = re.compile(r"^###\s*(.+?)\s*$", re.MULTILINE)
    matches = list(pattern.finditer(body))
    if not matches:
        return [("", body)]
    out = []
    for i, m in enumerate(matches):
        title = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(body)
        out.append((title, body[start:end].strip()))
    return out


def parse_dd_table(body):
    """Parse a markdown pipe-table into a list of dict rows (Data Dictionary)."""
    lines = [l.strip() for l in body.splitlines() if l.strip().startswith("|")]
    if len(lines) < 2:
        return []
    header = [c.strip().lower() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:  # skip header + separator row
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) != len(header):
            continue
        row = dict(zip(header, cells))
        rows.append(row)
    return rows


def render_mermaid_to_png(mermaid_code, out_path):
    mmd_path = out_path.with_suffix(".mmd")
    mmd_path.write_text(mermaid_code, encoding="utf-8")
    cmd = [
        "npx", "--yes", "@mermaid-js/mermaid-cli",
        "-i", str(mmd_path), "-o", str(out_path),
        "-b", "white", "-s", "2",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, shell=(sys.platform == "win32"))
    if result.returncode != 0 or not out_path.exists():
        raise RuntimeError(
            f"mermaid-cli failed for {out_path.name}:\n{result.stdout}\n{result.stderr}"
        )
    return out_path


def build_subdoc(tpl, text):
    """Turn prose (markdown-ish: blank-line paragraphs, '- ' bullets, **bold**) into a docxtpl subdoc."""
    sub = tpl.new_subdoc()
    bold_re = re.compile(r"\*\*(.+?)\*\*")
    # force a standalone bold "label" line (e.g. **Functional Requirements**) onto its own
    # block so it renders as a bold sub-heading instead of being run into the bullets below it
    text = re.sub(r"(?m)^(\*\*[^\n]+\*\*)\n(?=[-*] )", r"\1\n\n", text)
    for block in re.split(r"\n\s*\n", text.strip()):
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        is_list = all(l.strip().startswith(("- ", "* ")) for l in lines if l.strip())
        if is_list:
            for line in lines:
                item = line.strip().lstrip("-*").strip()
                p = sub.add_paragraph(style="List Bullet")
                _add_runs_with_bold(p, item, bold_re)
        else:
            p = sub.add_paragraph()
            _add_runs_with_bold(p, " ".join(l.strip() for l in lines), bold_re)
    if not sub.paragraphs:
        sub.add_paragraph("")
    return sub


def _add_runs_with_bold(paragraph, text, bold_re):
    pos = 0
    for m in bold_re.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("sdd_markdown", type=Path)
    ap.add_argument("output_docx", type=Path)
    ap.add_argument("--meta", type=Path, default=None, help="JSON with project_name/client_name/doc_version/doc_date/author/revisions")
    args = ap.parse_args()

    md_text = args.sdd_markdown.read_text(encoding="utf-8")
    sections = split_sections(md_text)

    meta = {
        "project_name": "ชื่อโครงการ (Project Name)",
        "client_name": "ชื่อลูกค้า (Client Name)",
        "doc_version": "1.0",
        "doc_date": "",
        "author": "",
        "revisions": [{"version": "1.0", "date": "", "author": "", "description": "First release"}],
    }
    if args.meta and args.meta.exists():
        meta.update(json.loads(args.meta.read_text(encoding="utf-8")))

    diagrams_dir = args.output_docx.with_name(args.output_docx.stem + "_diagrams")
    diagrams_dir.mkdir(parents=True, exist_ok=True)

    tpl = DocxTemplate(str(TEMPLATE_PATH))

    ctx = dict(meta)

    ctx["s1_intro"] = build_subdoc(tpl, extract_mermaid_blocks(sections.get(1, ""))[0])
    ctx["s2_requirements"] = build_subdoc(tpl, extract_mermaid_blocks(sections.get(2, ""))[0])
    ctx["s3_modules"] = build_subdoc(tpl, extract_mermaid_blocks(sections.get(3, ""))[0])
    ctx["s4_datamodel"] = build_subdoc(tpl, extract_mermaid_blocks(sections.get(4, ""))[0])

    # Section 5 — DFD: expect an L0 subheading then an L1 subheading, each with one mermaid block
    dfd_subs = extract_subheadings(sections.get(5, ""))
    l0_png = l1_png = None
    for title, body in dfd_subs:
        _, blocks = extract_mermaid_blocks(body)
        if not blocks:
            continue
        target = diagrams_dir / f"dfd_{'l0' if l0_png is None else 'l1'}.png"
        render_mermaid_to_png(blocks[0], target)
        if l0_png is None:
            l0_png = target
        else:
            l1_png = target
    ctx["s5_dfd_l0_img"] = InlineImage(tpl, str(l0_png), width=Mm(150)) if l0_png else ""
    ctx["s5_dfd_l1_img"] = InlineImage(tpl, str(l1_png), width=Mm(150)) if l1_png else ""

    # Section 6 — Flow Diagrams: one subheading + mermaid block per use case
    flows = []
    for idx, (title, body) in enumerate(extract_subheadings(sections.get(6, "")), start=1):
        _, blocks = extract_mermaid_blocks(body)
        if not blocks:
            continue
        target = diagrams_dir / f"flow_{idx}.png"
        render_mermaid_to_png(blocks[0], target)
        flows.append({"title": title or f"Flow {idx}", "image": InlineImage(tpl, str(target), width=Mm(150))})
    ctx["s6_flows"] = flows

    # Section 7 — ER Diagram
    _, er_blocks = extract_mermaid_blocks(sections.get(7, ""))
    if er_blocks:
        er_png = diagrams_dir / "er_diagram.png"
        render_mermaid_to_png(er_blocks[0], er_png)
        ctx["s7_er_img"] = InlineImage(tpl, str(er_png), width=Mm(160))
    else:
        ctx["s7_er_img"] = ""

    # Section 8 — Data Dictionary
    dd_rows = parse_dd_table(sections.get(8, ""))
    ctx["dd_rows"] = [
        {
            "entity": r.get("entity", ""),
            "field": r.get("field", ""),
            "type": r.get("type", ""),
            "null": r.get("null", ""),
            "key": r.get("key", ""),
            "constraint": r.get("constraint", ""),
            "description": r.get("description", ""),
        }
        for r in dd_rows
    ]
    if not ctx["dd_rows"]:
        ctx["dd_rows"] = [{"entity": "-", "field": "-", "type": "-", "null": "-", "key": "-", "constraint": "-", "description": "no data dictionary rows found"}]

    # Section 9 — Sitemap
    _, sitemap_blocks = extract_mermaid_blocks(sections.get(9, ""))
    if sitemap_blocks:
        sitemap_png = diagrams_dir / "sitemap.png"
        render_mermaid_to_png(sitemap_blocks[0], sitemap_png)
        ctx["s9_sitemap_img"] = InlineImage(tpl, str(sitemap_png), width=Mm(150))
    else:
        ctx["s9_sitemap_img"] = ""

    # Section 10 — User Roles & Permissions
    role_rows = parse_dd_table(sections.get(10, ""))
    ctx["role_rows"] = [
        {
            "name": r.get("role", r.get("name", "")),
            "permissions": r.get("permissions", ""),
            "pages": r.get("pages", r.get("accessible pages", "")),
        }
        for r in role_rows
    ]
    if not ctx["role_rows"]:
        ctx["role_rows"] = [{"name": "-", "permissions": "-", "pages": "no role table found"}]

    tpl.render(ctx)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(args.output_docx))
    print(f"Wrote {args.output_docx}")


if __name__ == "__main__":
    main()
